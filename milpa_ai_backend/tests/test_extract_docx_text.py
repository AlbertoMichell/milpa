"""Tests de extracción para DOCX y TXT/MD con paginación lógica.

Cubre las garantías que el endpoint ``/api/documents/ingest`` necesita:
  - DOCX preserva tablas y encabezados como bloques distinguibles.
  - DOCX asigna páginas lógicas según saltos de sección o palabras.
  - TXT/MD detecta form-feed (``\\f``) como salto de página real.
  - Markdown promueve h1 a inicio de página solo cuando NO hay marcador.
"""
from __future__ import annotations

import os
import tempfile
from pathlib import Path

import pytest


def _have_docx() -> bool:
    try:
        import docx  # noqa: F401
        return True
    except Exception:
        return False


@pytest.mark.skipif(not _have_docx(), reason="python-docx no instalado")
def test_docx_pages_and_tables(tmp_path: Path) -> None:
    """Genera DOCX con tabla y page-break, valida la estructura extraída."""
    from docx import Document
    from docx.enum.text import WD_BREAK
    from core.logic.extract_docx import extract_docx

    doc = Document()
    doc.add_heading("Manual prueba", level=1)
    doc.add_paragraph("Introducción al documento sintético.")
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)
    doc.add_heading("Sección 2", level=2)
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "A"
    table.rows[0].cells[1].text = "B"
    table.add_row().cells[0].text = "1"
    table.rows[1].cells[1].text = "2"

    docx_path = tmp_path / "t.docx"
    doc.save(docx_path)

    e = extract_docx(str(docx_path))
    assert e.n_pages >= 2, "el page-break explícito debe contar"
    kinds = {b.kind for b in e.blocks}
    assert "heading" in kinds and "table" in kinds
    tables = e.get_tables()
    assert len(tables) == 1
    assert tables[0].table_rows is not None
    assert ["A", "B"] in tables[0].table_rows


def test_text_form_feed_paginates(tmp_path: Path) -> None:
    """Form feed ``\\f`` debe dividir páginas; sin marcadores cae a paginación
    por palabras."""
    from core.logic.extract_text import extract_text

    p = tmp_path / "t.txt"
    p.write_text("Página uno línea uno\nlínea dos\f\nPágina dos contenido\n", encoding="utf-8")
    e = extract_text(str(p))
    pages = e.texts_per_page()
    assert len(pages) == 2
    assert "Página uno" in pages[1]
    assert "Página dos" in pages[2]


def test_markdown_h1_pagination(tmp_path: Path) -> None:
    """Sin marcadores de form-feed, los h1 promueven nueva página."""
    from core.logic.extract_text import extract_text

    p = tmp_path / "t.md"
    p.write_text(
        "# Capítulo uno\n\nContenido del primer capítulo.\n\n"
        "## Subsección\n\nDetalle.\n\n"
        "# Capítulo dos\n\nContenido distinto.\n",
        encoding="utf-8",
    )
    e = extract_text(str(p))
    assert e.n_pages == 2, f"esperado 2 páginas, fue {e.n_pages}"
    pages = e.texts_per_page()
    assert "Capítulo uno" in pages[1]
    assert "Capítulo dos" in pages[2]


def test_markdown_form_feed_takes_priority(tmp_path: Path) -> None:
    """Si hay form-feed, los h1 no añaden páginas extra (priorizamos lo
    explícito)."""
    from core.logic.extract_text import extract_text

    p = tmp_path / "t.md"
    p.write_text(
        "# Uno\n\nContenido uno.\n\f\n# Dos\n\nContenido dos.\n",
        encoding="utf-8",
    )
    e = extract_text(str(p))
    assert e.n_pages == 2


def test_text_default_paginates_long_text(tmp_path: Path) -> None:
    """Texto sin marcadores se pagina cada ~500 palabras."""
    from core.logic.extract_text import extract_text

    word = "palabra "
    long_text = (word * 600) + "\n\n" + (word * 600)
    p = tmp_path / "t.txt"
    p.write_text(long_text, encoding="utf-8")
    e = extract_text(str(p))
    assert e.n_pages >= 2
