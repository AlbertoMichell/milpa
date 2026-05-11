"""Genera un DOCX de prueba con encabezados, listas, tablas y saltos de página.

Uso:
    py -3 milpa_ai_backend/tools/gen_stress_docx.py --out path.docx
"""
from __future__ import annotations

import argparse
from pathlib import Path


def build_docx(out_path: Path) -> None:
    from docx import Document
    from docx.enum.text import WD_BREAK

    doc = Document()
    doc.add_heading("Manual MILPA — Prueba DOCX sintético", level=1)

    doc.add_paragraph(
        "Este documento ejercita la extracción DOCX preservando estructura: "
        "encabezados, listas, tablas y saltos de página explícitos. Marca "
        "canario uno: ROBLE_DOCX_INTRO_CANARIO_AB1."
    )

    doc.add_heading("1. Sección con lista", level=2)
    bullets = [
        "Maíz criollo de altura — selección por color y tamaño de mazorca.",
        "Frijol negro — fijador de nitrógeno asociado al maíz.",
        "Calabaza pipiana — cobertor del suelo, controla malezas.",
        "Marca canario lista: ALAMO_DOCX_LISTA_CANARIO_CD2.",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)

    doc.add_heading("2. Tabla de rendimientos", level=2)
    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = "Cultivo"
    hdr[1].text = "Variedad"
    hdr[2].text = "Rendimiento (t/ha)"
    hdr[3].text = "Notas"
    rows = [
        ("Maíz", "Cónico criollo", "3.2", "Marca canario tabla: ENCINA_DOCX_TABLA_CANARIO_EF3"),
        ("Frijol", "Negro Querétaro", "1.4", "Asociado al maíz"),
        ("Calabaza", "Pipiana", "8.0", "Pulpa para semilla"),
        ("Chile", "Serrano", "12.5", "Insumo para conservas"),
    ]
    for cult, var, rdmt, notes in rows:
        cells = table.add_row().cells
        cells[0].text = cult
        cells[1].text = var
        cells[2].text = rdmt
        cells[3].text = notes

    doc.add_paragraph(
        "Cierre del bloque tabular: la asociación de cultivos mejora el rendimiento "
        "agregado y reduce el uso de insumos sintéticos."
    )

    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)

    doc.add_heading("3. Recomendaciones de manejo", level=2)
    doc.add_paragraph(
        "El manejo agroecológico contempla rotación, cobertura del suelo y "
        "monitoreo de plagas. Marca canario manejo: PINO_DOCX_MANEJO_CANARIO_GH4. "
        "El sistema MILPA documenta cada práctica con su evidencia."
    )
    numbered = [
        "Diagnóstico inicial del lote (suelo, agua y antecedentes).",
        "Diseño de rotación con leguminosas y cucurbitáceas.",
        "Aplicación de compost y mulch antes de siembra.",
        "Monitoreo semanal de plagas con trampas amarillas.",
    ]
    for n in numbered:
        doc.add_paragraph(n, style="List Number")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--out",
        default="milpa_ai_backend/data/stress_docx_milpa.docx",
        help="Ruta de salida del DOCX",
    )
    args = parser.parse_args()
    out = Path(args.out)
    build_docx(out)
    print(out.resolve())


if __name__ == "__main__":
    main()
