# milpa_ai_backend/core/logic/extraction.py
# ------------------------------------------------------------
# DEPRECATED — pipeline alternativo no cableado al ingest activo.
#
# El pipeline operativo vive en `core/logic/extract.py` (extract_document para
# PDF, extract_docx_to_db para DOCX y extract_text_to_db para TXT/MD). Este
# módulo retiene utilidades de extracción de bbox por span para usos futuros
# (visor PDF clic-through). NO IMPORTAR desde el endpoint /api/documents/ingest
# — dejaría dos pipelines coexistiendo.
# ------------------------------------------------------------
from __future__ import annotations

from dataclasses import dataclass, field, asdict
from typing import List, Dict, Optional, Tuple, Any
import os

# Dependencias opcionales; producimos errores claros si faltan
try:
    import fitz  # PyMuPDF
except Exception as e:  # pragma: no cover
    fitz = None

try:
    import docx  # python-docx
except Exception:
    docx = None


@dataclass
class Span:
    text: str
    bbox: Optional[Tuple[float, float, float, float]]  # (x1,y1,x2,y2)
    font: Optional[str] = None
    size: Optional[float] = None


@dataclass
class PageExtraction:
    page: int
    text: str
    spans: List[Span] = field(default_factory=list)
    dsi: Optional[float] = None
    coverage: Optional[float] = None
    tables: List[Dict[str, Any]] = field(default_factory=list)   # llenará tables.py
    figures: List[Dict[str, Any]] = field(default_factory=list)  # [{'page', 'bbox', 'caption'}]


@dataclass
class ExtractionResult:
    mimetype: str
    pages: List[PageExtraction]
    doc_meta: Dict[str, Any] = field(default_factory=dict)


# -------------------------
# Heurísticas de métricas
# -------------------------
def _heuristic_coverage_by_spans(spans: List[Span]) -> float:
    """
    Heurística simple: cobertura ~ (#caracteres visibles) / (#spans * 50)
    (asumiendo 50 chars promedio por span). Cap a [0, 1].
    *No representa área real, pero correlaciona con densidad de texto.*
    """
    if not spans:
        return 0.0
    total_chars = sum(len(s.text.strip()) for s in spans if s.text)
    denom = max(len(spans) * 50, 1)
    cov = min(total_chars / denom, 1.0)
    return round(cov, 4)


def _heuristic_dsi(blocks: List[Dict[str, Any]]) -> float:
    """
    DSI% (Document Structure Integrity) heurístico:
    Promedia señales: orden de lectura estable + proporción de bloques de texto vs imágenes.
    """
    if not blocks:
        return 0.0
    text_blocks = sum(1 for b in blocks if b.get("type") == 0)
    img_blocks = sum(1 for b in blocks if b.get("type") == 1)
    ratio_text = text_blocks / max(text_blocks + img_blocks, 1)
    # Suma penalización mínima si hay demasiadas imágenes sin texto.
    dsi = 0.7 * ratio_text + 0.3  # baseline 0.3 + peso por texto
    return round(min(max(dsi, 0.0), 1.0), 4)


# -------------------------
# Extracción por tipo
# -------------------------
def extract_pdf(path: str) -> ExtractionResult:
    """
    Extrae texto/spans/bboxes/figuras de un PDF usando PyMuPDF.
    Adicionalmente, llena dsi y coverage por página.
    """
    if fitz is None:
        raise RuntimeError("PyMuPDF (pymupdf) no está instalado. Instálalo para procesar PDF.")

    doc = fitz.open(path)
    pages: List[PageExtraction] = []
    all_figures: List[Dict[str, Any]] = []

    for i, p in enumerate(doc):
        page_num = i + 1
        # 'dict' expone bloques/lines/spans con bbox, 'rawdict' es aún más bajo nivel
        p_dict = p.get_text("dict")
        blocks = p_dict.get("blocks", []) or []
        spans: List[Span] = []
        text_accum: List[str] = []

        # Recorremos bloques
        for b in blocks:
            if b.get("type") == 0:  # texto
                for l in b.get("lines", []):
                    for s in l.get("spans", []):
                        s_text = s.get("text", "") or ""
                        s_bbox = tuple(s.get("bbox")) if s.get("bbox") else None
                        s_font = s.get("font")
                        s_size = s.get("size")
                        spans.append(Span(text=s_text, bbox=s_bbox, font=s_font, size=s_size))
                        text_accum.append(s_text)
            elif b.get("type") == 1:  # imagen -> figura candidata
                fig_bbox = tuple(b.get("bbox")) if b.get("bbox") else None
                all_figures.append({"page": page_num, "bbox": fig_bbox, "caption": None})

        # Métricas
        coverage = _heuristic_coverage_by_spans(spans)
        dsi = _heuristic_dsi(blocks)

        pages.append(
            PageExtraction(
                page=page_num,
                text="\n".join([t for t in text_accum if t is not None]),
                spans=spans,
                dsi=dsi,
                coverage=coverage,
                tables=[],   # tables.py llenará si lo conectas
                figures=[f for f in all_figures if f["page"] == page_num],
            )
        )

    return ExtractionResult(mimetype="application/pdf", pages=pages, doc_meta={"pages": len(pages)})


def extract_docx(path: str) -> ExtractionResult:
    """
    Extrae texto de DOCX con estructura básica (encabezados/listas), sin bbox.
    Nota: DOCX no ofrece coords; las páginas las consideramos 1 única (page=1).
    """
    if docx is None:
        raise RuntimeError("python-docx no está instalado. Instálalo para procesar DOCX.")

    d = docx.Document(path)
    lines: List[str] = []
    for p in d.paragraphs:
        lines.append(p.text)

    # Tablas DOCX (volcadas como CSV inline por simplicidad de MVP de extracción)
    for t in d.tables:
        for r in t.rows:
            row = [c.text.replace("\n", " ").strip() for c in r.cells]
            lines.append(" | ".join(row))

    page = PageExtraction(page=1, text="\n".join(lines), spans=[], dsi=0.95, coverage=1.0, tables=[], figures=[])
    return ExtractionResult(mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document", pages=[page])


def extract_txt(path: str) -> ExtractionResult:
    """
    Extrae texto plano en UTF-8 (con fallback). Bboxes no aplican.
    """
    # Solo usar UTF-8, sin fallback problemáticos
    try:
        with open(path, "r", encoding="utf-8") as f:
            content = f.read()
    except UnicodeDecodeError as e:
        # Si falla UTF-8, intentar UTF-8-SIG
        try:
            with open(path, "r", encoding="utf-8-sig") as f:
                content = f.read()
        except Exception:
            # Último recurso: leer como bytes y decodificar con reemplazo
            with open(path, "rb") as f:
                content = f.read().decode("utf-8", errors="replace")

    page = PageExtraction(page=1, text=content, spans=[], dsi=1.0, coverage=1.0, tables=[], figures=[])
    return ExtractionResult(mimetype="text/plain", pages=[page])


# -------------------------
# Facade según mimetype/ext
# -------------------------
def extract_document(path: str) -> ExtractionResult:
    """
    Facade: escoge extractor por extensión si el mimetype aún no está disponible.
    """
    ext = os.path.splitext(path.lower())[-1]
    if ext == ".pdf":
        return extract_pdf(path)
    if ext in (".docx",):
        return extract_docx(path)
    if ext in (".txt",):
        return extract_txt(path)
    # Por seguridad, intenta PDF y deja error claro si falla.
    return extract_pdf(path)


# -------------------------
# Fine refs (clic-through)
# -------------------------
def build_fine_ref_from_spans(page: PageExtraction) -> Optional[Tuple[int, Tuple[float, float, float, float]]]:
    """
    Selecciona un bbox representativo para la página:
    - Si hay spans: toma el bbox del span medio.
    - Si no hay spans: None (DOCX/TXT típicamente).
    Retorna (page_number, bbox).
    """
    if not page.spans:
        return None
    mid = page.spans[len(page.spans) // 2]
    if mid.bbox is None:
        return None
    return (page.page, mid.bbox)


def to_serializable(result: ExtractionResult) -> Dict[str, Any]:
    """Convierte dataclasses a dicts puros (útil para debug/logs)."""
    data = {
        "mimetype": result.mimetype,
        "doc_meta": result.doc_meta,
        "pages": [],
    }
    for p in result.pages:
        data["pages"].append(
            {
                "page": p.page,
                "text": p.text,
                "dsi": p.dsi,
                "coverage": p.coverage,
                "spans": [asdict(s) for s in p.spans],
                "tables": p.tables,
                "figures": p.figures,
            }
        )
    return data
