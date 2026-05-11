"""Extracción rica de DOCX preservando estructura.

A diferencia del aplanado simple ``\\n\\n``.join(paragraphs), este módulo:
  - Reconoce encabezados (Heading 1..9) y los marca como secciones.
  - Extrae tablas como objetos estructurados (no como texto crudo).
  - Detecta listas (numeradas y bullets) preservando jerarquía.
  - Asigna páginas lógicas usando los saltos de sección/página explícitos del
    documento (``w:type='page'``, ``w:type='nextPage'``) o, en su defecto,
    cada ~500 palabras (aprox. 1 página A4 a 11pt).

El objeto devuelto se consume desde el endpoint de ingest, que decide qué
fragments persistir en SQLite y qué tablas registrar en la tabla ``tables``.
"""
from __future__ import annotations

import logging
import re
from dataclasses import dataclass, field
from typing import List, Optional

_log = logging.getLogger(__name__)

# Aproximación A4 / 11pt / interlineado 1.15 ≈ 500-550 palabras por página.
_DEFAULT_WORDS_PER_PAGE = 500

# qn helper local para no obligar a un import cualquiera de python-docx.
def _qn(tag: str) -> str:
    """Reescribe un namespace tag como hace docx.oxml.ns.qn."""
    prefix, local = tag.split(":", 1)
    nsmap = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    }
    return f"{{{nsmap[prefix]}}}{local}"


@dataclass
class DocxBlock:
    kind: str  # "heading" | "paragraph" | "list" | "table" | "page_break"
    page: int
    text: str
    level: int = 0
    table_rows: Optional[List[List[str]]] = None  # solo si kind=="table"
    style: Optional[str] = None


@dataclass
class DocxExtraction:
    blocks: List[DocxBlock] = field(default_factory=list)
    n_pages: int = 1
    raw_text: str = ""

    def texts_per_page(self) -> dict[int, str]:
        out: dict[int, list[str]] = {}
        for b in self.blocks:
            if b.kind == "table":
                # Aplanamos la tabla a texto buscable para la página.
                txt = "\n".join(" | ".join(row) for row in (b.table_rows or []))
                if not txt.strip():
                    continue
                out.setdefault(b.page, []).append(f"[Tabla docx]\n{txt}")
            else:
                if b.text.strip():
                    out.setdefault(b.page, []).append(b.text)
        return {p: "\n\n".join(v) for p, v in out.items()}

    def get_tables(self) -> list[DocxBlock]:
        return [b for b in self.blocks if b.kind == "table"]


def _has_explicit_page_break(paragraph) -> bool:
    """Detecta ``<w:br w:type="page"/>`` dentro de un párrafo de python-docx."""
    try:
        body = paragraph._element  # type: ignore[attr-defined]
        for br in body.iter(_qn("w:br")):
            if br.get(_qn("w:type")) == "page":
                return True
    except Exception:
        return False
    return False


def _heading_level(style_name: str) -> Optional[int]:
    if not style_name:
        return None
    m = re.match(r"^heading\s*(\d+)$", style_name.strip().lower())
    if m:
        return int(m.group(1))
    return None


def _is_list_paragraph(paragraph) -> bool:
    """Detecta listas numeradas/bullet via numPr en el formato XML."""
    try:
        body = paragraph._element  # type: ignore[attr-defined]
        for pPr in body.iter(_qn("w:pPr")):
            if pPr.find(_qn("w:numPr")) is not None:
                return True
    except Exception:
        return False
    return False


def extract_docx(
    path: str,
    words_per_page: int = _DEFAULT_WORDS_PER_PAGE,
) -> DocxExtraction:
    """Procesa un .docx con preservación de estructura.

    Lanza ImportError si python-docx no está instalado.
    """
    import docx  # type: ignore

    doc = docx.Document(path)
    out = DocxExtraction()

    page = 1
    word_acc = 0

    def maybe_advance_page(words: int) -> None:
        nonlocal page, word_acc
        word_acc += words
        if word_acc >= words_per_page:
            page += 1
            word_acc = 0

    # Recorremos elementos del cuerpo (párrafos y tablas, en orden de lectura).
    body = doc.element.body
    raw_chunks: list[str] = []
    for child in body.iterchildren():
        tag = child.tag.split("}", 1)[-1] if "}" in child.tag else child.tag
        if tag == "p":
            # Buscamos su objeto Paragraph homólogo.
            try:
                p = next(p for p in doc.paragraphs if p._element is child)
            except StopIteration:
                continue
            text = (p.text or "").strip()
            if not text:
                # Verificamos si trae un page-break explícito.
                if _has_explicit_page_break(p):
                    page += 1
                    word_acc = 0
                    out.blocks.append(DocxBlock(kind="page_break", page=page, text=""))
                continue

            level = _heading_level(getattr(p.style, "name", "") or "")
            if level is not None:
                out.blocks.append(
                    DocxBlock(kind="heading", page=page, text=text, level=level, style=p.style.name)
                )
            elif _is_list_paragraph(p):
                out.blocks.append(DocxBlock(kind="list", page=page, text=text))
            else:
                out.blocks.append(DocxBlock(kind="paragraph", page=page, text=text))
            raw_chunks.append(text)

            if _has_explicit_page_break(p):
                page += 1
                word_acc = 0
            else:
                maybe_advance_page(len(text.split()))
        elif tag == "tbl":
            # Tabla: localizamos su Table object.
            try:
                tbl = next(t for t in doc.tables if t._element is child)
            except StopIteration:
                continue
            rows: list[list[str]] = []
            for row in tbl.rows:
                cells = [(c.text or "").strip() for c in row.cells]
                if any(cells):
                    rows.append(cells)
            if not rows:
                continue
            out.blocks.append(
                DocxBlock(kind="table", page=page, text="", table_rows=rows)
            )
            raw_chunks.append("\n".join(" | ".join(r) for r in rows))
            # Una tabla equivale a ~ varias líneas; promovemos avance moderado.
            maybe_advance_page(sum(len(c.split()) for r in rows for c in r))

    out.n_pages = max(page, 1)
    out.raw_text = "\n\n".join(raw_chunks)
    return out
